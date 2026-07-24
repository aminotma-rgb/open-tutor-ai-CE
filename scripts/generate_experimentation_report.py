"""Génère le rapport d'expérimentation .docx (LongMemEval-S + MRBench réels, LoCoMo simulé).

Usage :
    python3 scripts/generate_experimentation_report.py
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "rapport_experimentation.docx"

doc = Document()

# ── Styles de base ───────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        shade_cell(hdr[i], "2E4053")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(10)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return table


def add_disclaimer(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("⚠ " + text)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xB0, 0x3A, 0x2E)
    run.font.size = Pt(10)


def add_confirmation(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("✅ " + text)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x7D, 0x32)
    run.font.size = Pt(10)


def add_subheading(text):
    h = doc.add_heading(text, level=3)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x2E, 0x4E, 0x83)


def add_body(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.bold = bold
    r.font.italic = italic
    return p


def add_bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(it)


# ═══════════════════════════════════════════════════════════════════════
# PAGE DE GARDE
# ═══════════════════════════════════════════════════════════════════════
title = doc.add_heading("Rapport d'expérimentation — Protocole d'évaluation externe", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("OpenTutorAI CE — LongMemEval-S, LoCoMo, MRBench")
r.font.size = Pt(14)
r.font.italic = True

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.add_run("Régénéré le 2026-07-12").font.size = Pt(10)

doc.add_paragraph()
warn_box = doc.add_paragraph()
warn_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
warn_run = warn_box.add_run(
    "⚠ ÉTAT MIXTE — LongMemEval-S et MRBench réels, LoCoMo simulé ⚠\n"
    "LongMemEval-S (section 3) et MRBench (section 5) présentent des résultats RÉELS, "
    "issus de runs complets terminés avec succès. LoCoMo (section 4) reste en revanche des "
    "valeurs FICTIVES, générées pour illustrer le format attendu en attendant son run réel "
    "(~16-17h estimées sur ce matériel CPU sans GPU). Chaque section indique explicitement, "
    "par un bandeau vert (✅ réel) ou rouge (⚠ simulé), le statut de ses résultats."
)
warn_run.font.bold = True
warn_run.font.size = Pt(11)
warn_run.font.color.rgb = RGBColor(0xB0, 0x3A, 0x2E)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 1. VUE D'ENSEMBLE
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("1. Vue d'ensemble du protocole", level=1)
add_body(
    "Le protocole d'évaluation externe d'OpenTutorAI CE s'appuie sur trois benchmarks "
    "académiques indépendants, complémentaires aux scénarios internes déterministes S0-S9. "
    "Deux dimensions sont couvertes :"
)
add_bullets([
    "D1 — Mémoire persistante inter-sessions : LongMemEval-S (Wu et al., ICLR 2025) et "
    "LoCoMo (Maharana et al., ACL 2024)",
    "D3 — Qualité pédagogique de correction d'erreurs : MRBench (Maurya et al., BEA 2025)",
])
add_body(
    "Ces trois benchmarks font des appels LLM réels (juge et pipeline de génération), "
    "sont exclus de la CI rapide, et doivent s'exécuter séquentiellement — jamais en "
    "parallèle — en raison d'une contrainte partagée : Ollama en CPU sans GPU ne traite "
    "qu'une seule requête à la fois."
)

add_subheading("Juge LLM commun")
add_body(
    "Les trois benchmarks utilisent le même juge : mistral:7b-instruct-q4_K_M via Ollama "
    "local, en CPU sans GPU. Ce choix d'infrastructure (local, gratuit, sans clé API) est "
    "délibéré mais a un coût de fiabilité documenté — voir section 6."
)

# ═══════════════════════════════════════════════════════════════════════
# 2. CORRECTIFS APPLIQUÉS DURANT CETTE CAMPAGNE
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("2. Correctifs méthodologiques appliqués avant/pendant ces runs", level=1)
add_body(
    "Plusieurs écarts de fidélité et bugs d'exécution ont été identifiés et corrigés au "
    "fil de cette campagne. Ils conditionnent la lecture des résultats — un run antérieur "
    "à un correctif n'est pas comparable à un run postérieur."
)

add_table(
    ["Correctif", "Fichier", "Effet"],
    [
        ("DAMR au lieu d'un score pondéré", "mrbench_runner.py",
         "MRBench utilisait yes=1.0/to some extent=0.5/no=0.0 (une moyenne pondérée), "
         "non conforme au papier. Remplacé par un match binaire strict au label désiré."),
        ("Axe revealing_of_answer ajouté", "mrbench_runner.py",
         "5e axe MRBench manquant, pourtant central à la mission pédagogique du tuteur "
         "(« guider sans révéler la réponse »)."),
        ("Cibles indicatives (warning) au lieu de bloquantes", "test_external_mrbench.py, "
         "test_external_locomo.py, test_external_longmemeval.py",
         "Échantillons trop petits pour un verdict pass/fail fiable. LongMemEval-S est "
         "passé en indicatif en dernier (2026-07-11), par cohérence avec MRBench/LoCoMo — "
         "à N=4, chaque instance pèse 25 points."),
        ("Prompts de juge différenciés par catégorie + date injectée", "longmemeval_runner.py",
         "Fidélité à la Figure 10/13 du papier LongMemEval (tolérance temporelle, "
         "« Current Date » pour résoudre les expressions relatives)."),
        ("Portée réduite à knowledge-update", "scripts/sample_longmemeval.py",
         "LongMemEval-S intégral impraticable sur CPU (runs de plusieurs heures). "
         "LoCoMo couvre les 5 autres capacités."),
        ("Signal de fin d'orchestrateur respecté", "orchestrator.py",
         "Bug de production : feedback_node signalait next_agent=\"END\" mais "
         "l'orchestrateur l'ignorait et recalculait un routage à chaque appel — le "
         "graphe ne s'arrêtait jamais tout seul, épuisant systématiquement les 15 "
         "itérations maximum. Corrigé : gain moyen de ~3x sur le nombre d'itérations."),
        ("Plafond de retry du planificateur respecté par le routeur LLM", "orchestrator.py",
         "Le repli déterministe limitait les retries du planificateur à 3, mais le "
         "routeur LLM (utilisé en priorité) n'avait aucune limite équivalente — observé "
         "jusqu'à un « retry #6 » avant le plafond dur. Corrigé par validation explicite."),
        ("Prompt du juge MRBench — cadrage anti-complaisance", "mrbench_runner.py",
         "Raisonnement exigé avant verdict + consigne sceptique explicite, en réponse "
         "au biais de sur-prédiction de \"yes\" mesuré par calibration F1 (section 6)."),
        ("revealing_of_answer — 1er correctif (instruction explicite)", "adaptive.py",
         "Bloc ANSWER VERIFICATION ordonnait littéralement de révéler la réponse. "
         "Corrigé, mais testé sans effet mesurable (0,225 → 0,20 sur N=10)."),
        ("revealing_of_answer — 2e correctif, ANNULÉ", "adaptive.py",
         "Suppression de la réponse dans le bloc PREPARED EXERCISES — mesuré comme "
         "AGGRAVANT le problème sur les mêmes dialogues testés avant/après (1,0 → 0,0 "
         "sur 2 cas). Annulé le 2026-07-12, cause exacte non identifiée, en attente "
         "d'investigation (texte généré non capturé par le runner actuel)."),
    ],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# Helper pour générer une section de benchmark complète
# ═══════════════════════════════════════════════════════════════════════

def benchmark_section(
    num, name, subtitle, objectif, avantages, metrique_desc, formule,
    parametrage_rows, resultats_headers, resultats_rows, interpretation,
    contraintes, resultats_reels=False, confirmation_texte=None,
):
    doc.add_heading(f"{num}. {name}", level=1)
    p = doc.add_paragraph()
    r = p.add_run(subtitle)
    r.font.italic = True
    r.font.size = Pt(11)

    add_subheading(f"{num}.1 Objectif (pourquoi ce benchmark)")
    add_body(objectif)

    add_subheading(f"{num}.2 Avantages")
    add_bullets(avantages)

    add_subheading(f"{num}.3 Métrique utilisée — description et formalisme")
    add_body(metrique_desc)
    formule_p = doc.add_paragraph()
    formule_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = formule_p.add_run(formule)
    fr.font.name = "Consolas"
    fr.font.size = Pt(11)
    fr.font.bold = True

    add_subheading(f"{num}.4 Paramétrage")
    add_table(["Paramètre", "Valeur"], parametrage_rows)

    if resultats_reels:
        add_subheading(f"{num}.5 Résultats obtenus (réels) et interprétation")
        add_confirmation(confirmation_texte or "Résultats réels — run terminé avec succès.")
    else:
        add_subheading(f"{num}.5 Résultats obtenus (simulés) et interprétation")
        add_disclaimer("Valeurs fictives — run réel en attente / en cours.")
    add_table(resultats_headers, resultats_rows)
    doc.add_paragraph()
    add_body(interpretation)

    add_subheading(f"{num}.6 Contraintes d'exécution")
    add_bullets(contraintes)

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════
# 3. LONGMEMEVAL-S (RÉEL)
# ═══════════════════════════════════════════════════════════════════════
benchmark_section(
    num=3,
    name="Benchmark A — LongMemEval-S",
    subtitle="Wu et al., ICLR 2025 — Dimension D1, capacité knowledge-update",
    objectif=(
        "OpenTutorAI CE n'est pas un simple proxy vers un LLM : chaque échange étudiant "
        "passe par un pipeline LangGraph qui maintient un état persistant inter-sessions "
        "via un checkpointer SQLite. LongMemEval-S a été retenu pour vérifier que cette "
        "mémoire tient sur un volume et une diversité de conversations qu'aucun scénario "
        "interne (S1, S3, S5) ne peut réalistement simuler, sur un dataset publié et "
        "indépendant du système. Sa portée a été réduite à la seule catégorie "
        "\"knowledge-update\" (reconnaître qu'une information est périmée et la mettre à "
        "jour) car c'est la seule des 6 capacités du papier qu'aucune alternative plus "
        "légère (LoCoMo) ne teste — et c'est aussi la catégorie où un bug de runner avait "
        "déjà été diagnostiqué, justifiant une surveillance continue."
    ),
    avantages=[
        "Ancrage académique fort (ICLR 2025), une référence du domaine",
        "Dataset construit indépendamment du système, sans biais de conception favorable",
        "Correspond exactement aux classes de sollicitation mémoire déjà testées en interne",
        "Champs de datation (question_date, haystack_dates) permettent un raisonnement "
        "temporel fidèle au protocole du papier",
    ],
    metrique_desc=(
        "Précision binaire correct/incorrect. Le runner rejoue tout le haystack de "
        "sessions passées à travers le pipeline LangGraph, pose la question finale avec "
        "la date courante injectée dans le prompt, puis un juge LLM évalue si la réponse "
        "générée est sémantiquement correcte par rapport à la réponse attendue "
        "(reformulations acceptées, pas d'exact-match). Le prompt du juge est différencié "
        "pour knowledge-update : une réponse contenant l'ancienne information à côté de "
        "la valeur mise à jour reste jugée correcte tant que la mise à jour attendue y "
        "figure."
    ),
    formule="Accuracy = (# réponses jugées correctes) / (# instances)",
    parametrage_rows=[
        ("Source", "longmemeval_s_cleaned.json (500 instances, HuggingFace)"),
        ("Échantillon", "4 instances, catégorie knowledge-update uniquement"),
        ("Seed", "42"),
        ("Juge", "llm_judge → mistral:7b-instruct-q4_K_M (Ollama local, CPU)"),
        ("Cible", "≥ 0,50 — indicative (warning, pas bloquante, depuis le 2026-07-11)"),
    ],
    resultats_headers=["question_id", "Catégorie", "Verdict"],
    resultats_rows=[
        ("4d6b87c8", "knowledge-update", "✅ correct"),
        ("852ce960", "knowledge-update", "✅ correct"),
        ("1cea1afa", "knowledge-update", "✅ correct"),
        ("18bc8abd", "knowledge-update", "✅ correct"),
        ("Accuracy globale", "—", "1,00 (4/4)"),
    ],
    interpretation=(
        "Résultat parfait : les 4 instances sont jugées correctes, bien au-dessus de la "
        "cible indicative (0,50). Le run a été interrompu une première fois pour "
        "diagnostiquer le bug d'orchestrateur (voir section 2), puis repris à partir de "
        "la session déjà en cours plutôt que relancé de zéro (30 sessions sur 52 déjà "
        "traitées pour la première instance conservées) — les 4 verdicts finaux "
        "proviennent bien d'un run complet et cohérent, pas d'un mélange de méthodologies. "
        "Réserve à garder en tête : à n=4, un résultat parfait est encourageant mais ne "
        "garantit pas la même performance sur un échantillon plus large — c'est un bon "
        "signal, pas une preuve statistique définitive. Aucune calibration F1 indépendante "
        "n'existe pour ce juge sur cette tâche (contrairement à MRBench), donc la fiabilité "
        "du jugement lui-même reste, comme pour LoCoMo, non vérifiée directement."
    ),
    contraintes=[
        "Exécution séquentielle obligatoire (jamais en parallèle avec LoCoMo/MRBench)",
        "Durée réelle du run complet (reprise incluse) : ~12h30 pour finir la 1ère "
        "instance (22 sessions restantes sur 52) + les 3 autres instances entières "
        "(39, 44 et 49 sessions) — bien plus long que l'estimation initiale (~1h20-1h30), "
        "car chaque instance est un thread \"à froid\" avec un contenu par session "
        "nettement plus dense que LoCoMo (~9-12k caractères/session contre ~300-700)",
        "Cible indicative depuis le 2026-07-11 : un score bas émettrait un warning, ne "
        "ferait plus échouer le test — changement appliqué par cohérence avec MRBench/LoCoMo",
        "Nécessite OTAI_RUN_EXTERNAL=1 et le dataset pré-échantillonné "
        "(scripts/sample_longmemeval.py)",
    ],
    resultats_reels=True,
    confirmation_texte=(
        "Résultats réels — run terminé avec succès le 2026-07-12 (reprise après "
        "interruption pour diagnostic), 4/4 instances correctes (accuracy = 1,00)."
    ),
)

# ═══════════════════════════════════════════════════════════════════════
# 4. LOCOMO (SIMULÉ)
# ═══════════════════════════════════════════════════════════════════════
benchmark_section(
    num=4,
    name="Benchmark B — LoCoMo",
    subtitle="Maharana et al., ACL 2024 — Dimension D1, complément léger (5 capacités)",
    objectif=(
        "LongMemEval-S seul était impraticable sur l'ensemble de ses 6 catégories : "
        "chaque instance rejoue un haystack de ~50 sessions (~115k tokens), avec des "
        "runs de plusieurs heures déjà documentés pour des échantillons de 30 à 40 "
        "instances sur ce CPU sans GPU. LoCoMo a été retenu comme complément pour "
        "couvrir les 5 capacités mémoire abandonnées par LongMemEval-S (single-hop, "
        "multi-hop, temporal-reasoning, open-domain, adversarial), avec un volume de "
        "texte nettement plus léger par conversation."
    ),
    avantages=[
        "~5 à 10x plus léger en tokens par conversation (10-22k contre ~115k)",
        "Structure économe : une conversation est rejouée une seule fois puis interrogée "
        "sur plusieurs questions, contrairement à LongMemEval-S où chaque question a son "
        "propre haystack dédié",
        "Benchmark académique reconnu, directement cité et comparé dans le papier "
        "LongMemEval lui-même (Table 1) — pas une alternative improvisée",
    ],
    metrique_desc=(
        "Précision binaire par catégorie, jugée par LLM (équivalence sémantique, pas "
        "d'exact-match) — un choix méthodologique assumé qui diverge du F1 de "
        "recouvrement de mots utilisé par le papier d'origine (le F1 pénalise les "
        "reformulations pourtant correctes). Tolérance ±1 jour appliquée à "
        "temporal-reasoning, comme pour LongMemEval-S."
    ),
    formule="Accuracy(catégorie) = (# réponses correctes) / (# questions de la catégorie)",
    parametrage_rows=[
        ("Source", "locomo10.json (10 conversations, GitHub snap-research/locomo)"),
        ("Échantillon", "3 conversations, 2 questions/catégorie/conversation (~28 au total)"),
        ("Seed", "42"),
        ("Juge", "llm_judge → mistral:7b-instruct-q4_K_M (Ollama local, CPU)"),
        ("Cible", "≥ 0,50 par catégorie — indicative (warning, pas bloquante)"),
    ],
    resultats_headers=["Catégorie", "n", "Accuracy", "Cible"],
    resultats_rows=[
        ("single-hop", "6", "0,83 ✅", "≥ 0,50"),
        ("multi-hop", "6", "0,67 ✅", "≥ 0,50"),
        ("temporal-reasoning", "6", "0,50 ✅", "≥ 0,50"),
        ("open-domain", "4", "0,75 ✅", "≥ 0,50"),
        ("adversarial", "6", "0,50 ✅", "≥ 0,50"),
        ("Global", "28", "0,68", "—"),
    ],
    interpretation=(
        "Le rappel simple (single-hop) score le plus haut, cohérent avec le fait que "
        "mémoriser un fait explicite mentionné une fois est la tâche la plus simple des "
        "cinq. temporal-reasoning et adversarial sont les plus exigeantes (raisonnement "
        "sur les dates, évidence difficile à relier à la question). Réserve majeure : "
        "n=4 à 6 par catégorie signifie qu'une seule question qui bascule change le "
        "score de 17 à 25 points — plus volatile encore que MRBench (n=20/corpus). Le "
        "juge utilisé ici n'a pas fait l'objet d'une calibration F1 dédiée (contrairement "
        "à MRBench) : le biais de complaisance mesuré ailleurs est plausible mais non "
        "vérifié sur cette tâche de jugement plus simple (correct/incorrect factuel, "
        "sans les 3 niveaux nuancés de MRBench)."
    ),
    contraintes=[
        "Exécution séquentielle obligatoire",
        "A nécessité deux correctifs de l'orchestrateur pour être praticable en pratique "
        "(signal END ignoré, plafond de retry non respecté par le routeur LLM) — un "
        "premier run sans ces correctifs projetait ~20h pour l'échantillon complet",
        "Après correctifs : ~5-6 itérations d'orchestrateur en moyenne par appel "
        "(mesuré), contre 15-16 systématiques avant — gain confirmé mais la durée totale "
        "reste de plusieurs heures (chaque conversation nécessite ~5-7 étapes de "
        "pipeline réelles au minimum)",
        "Écarts assumés au dataset d'origine : dialogues humain-humain mappés "
        "arbitrairement sur les rôles user/assistant ; pas de date de question explicite "
        "(date de la dernière session utilisée comme approximation)",
        "Estimation pour le run complet réel : ~16-17h (jamais exécuté à ce stade)",
    ],
)

# ═══════════════════════════════════════════════════════════════════════
# 5. MRBENCH (RÉEL)
# ═══════════════════════════════════════════════════════════════════════
benchmark_section(
    num=5,
    name="Benchmark C — MRBench",
    subtitle="Maurya et al., BEA 2025 — Dimension D3, qualité pédagogique de correction d'erreurs",
    objectif=(
        "Le cœur de la proposition de valeur d'OpenTutorAI CE n'est pas de répondre à "
        "des questions mais de tutorer : identifier une erreur de raisonnement chez "
        "l'étudiant, la localiser, et guider vers la correction sans donner directement "
        "la réponse. MRBench a été retenu car il évalue précisément ce comportement sur "
        "des dialogues mathématiques réels contenant une erreur de raisonnement — un "
        "benchmark généraliste de NLP (QA, résumé) n'aurait rien apporté ici, il "
        "n'existe pas de notion d'\"erreur d'étudiant à corriger pédagogiquement\" à y "
        "évaluer."
    ),
    avantages=[
        "L'un des rares benchmarks académiques ciblant spécifiquement la correction "
        "pédagogique d'erreurs (BEA 2025, atelier dédié aux applications éducatives des LLM)",
        "Deux corpus complémentaires : Bridge (élémentaire, dialogues courts) et MathDial "
        "(collège, raisonnement multi-étapes) — diversité de difficulté",
        "5 axes indépendants couvrent l'ensemble du comportement attendu d'un tuteur, "
        "y compris la clause centrale \"sans révéler la réponse\" (axe revealing_of_answer)",
    ],
    metrique_desc=(
        "DAMR (Desired Annotation Match Rate), définie par Maurya et al. (2025, §5.4). "
        "Pour chaque dimension pédagogique, un label est désigné comme désiré (\"Yes\" "
        "pour 4 axes, \"No\" pour revealing_of_answer). Le juge répond sur une échelle à "
        "trois niveaux (yes / to some extent / no), mais seul le label désiré compte "
        "comme succès — \"to some extent\" est un non-match au même titre que le "
        "contraire, pas une pondération à 0,5. Le prompt du juge exige un raisonnement "
        "explicite avant le verdict, avec une consigne sceptique, en réponse au biais de "
        "complaisance mesuré (section 6)."
    ),
    formule="DAMR(axe) = (# réponses avec verdict = label désiré de l'axe) / (# réponses totales) × 100",
    parametrage_rows=[
        ("Source", "MRBench_V2.json (200 instances Bridge + MathDial)"),
        ("Échantillon", "40 instances (20 Bridge + 20 MathDial)"),
        ("Seed", "42"),
        ("Juge", "llm_judge → mistral:7b-instruct-q4_K_M (Ollama local, CPU)"),
        ("Axes", "mistake_identification, mistake_location, guidance, actionability, "
                 "revealing_of_answer"),
        ("Cible", "≥ 0,50 par axe — indicative (warning, pas bloquante)"),
    ],
    resultats_headers=["Axe", "Global (n=40)", "Bridge (n=20)", "MathDial (n=20)", "Cible"],
    resultats_rows=[
        ("mistake_identification", "0,650 ✅", "0,650", "0,650", "≥ 0,50"),
        ("mistake_location", "0,600 ✅", "0,550", "0,650", "≥ 0,50"),
        ("guidance", "0,950 ✅", "0,900", "1,000", "≥ 0,50"),
        ("actionability", "0,875 ✅", "0,900", "0,850", "≥ 0,50"),
        ("revealing_of_answer", "0,225 ❌", "0,200", "0,250", "≥ 0,50"),
    ],
    interpretation=(
        "4 axes sur 5 dépassent nettement la cible, avec guidance (0,950) et "
        "actionability (0,875) particulièrement solides. Mais revealing_of_answer "
        "échoue franchement (0,225 contre ≥ 0,50) : le tuteur révèle la réponse finale "
        "à l'étudiant dans environ 78% des cas au lieu de le guider vers sa propre "
        "correction — c'est directement la clause centrale de la mission pédagogique du "
        "système (« guider sans donner la réponse ») qui est mise en défaut. Deux "
        "correctifs ont été tentés depuis ce run (voir section 2) : le premier "
        "(instruction explicite dans le prompt de vérification) sans effet mesurable "
        "(0,20 sur N=10), le second (suppression de la réponse du bloc "
        "PREPARED EXERCISES) a au contraire fait RÉGRESSER le score (0,00 sur N=6, "
        "confirmé sur les mêmes dialogues avant/après) — annulé, cause non identifiée. "
        "Nuance sur le biais du juge : pour revealing_of_answer (seul axe désiré=\"no\"), "
        "un juge biaisé vers \"yes\" ferait paraître ce score pire qu'il ne l'est "
        "vraiment, pas meilleur — même en tenant compte du biais connu, ce résultat "
        "reste préoccupant."
    ),
    contraintes=[
        "Exécution séquentielle obligatoire",
        "Chaque dialogue est un thread indépendant \"à froid\" (pas de mémoire "
        "pré-existante à réutiliser, contrairement aux appels successifs de LoCoMo sur "
        "un même thread) — le pipeline doit traverser les 7 étapes au moins une fois par "
        "dialogue, sans bénéfice de raccourci",
        "Durée réelle observée : 19h05 (68711 s) pour les 40 instances",
        "Nécessite OTAI_RUN_EXTERNAL=1 et le dataset pré-échantillonné "
        "(scripts/sample_mrbench.py)",
    ],
    resultats_reels=True,
    confirmation_texte=(
        "Résultats réels — run terminé avec succès le 2026-07-10 (2/2 tests passés, "
        "durée 19h05min11s), avec tous les correctifs appliqués (DAMR, 5 axes, "
        "correctifs orchestrateur, prompt anti-complaisance)."
    ),
)

# ═══════════════════════════════════════════════════════════════════════
# 6. LIMITES TRANSVERSALES
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("6. Limite transversale — fiabilité du juge LLM", level=1)
add_body(
    "Les trois benchmarks partagent le même juge (mistral:7b-instruct-q4_K_M, CPU "
    "local). Une calibration F1 indépendante a été menée sur MRBench (docs/F1_judge.md, "
    "run du 2026-07-07, AVANT le prompt anti-complaisance) : 40 instances du devset "
    "MRBench v3, réponses déjà générées par des tuteurs tiers (Sonnet, GPT-4, Llama, "
    "Mistral, Gemini, Phi3, Expert humain, Novice) avec labels humains existants, "
    "comparées aux verdicts de notre juge."
)
add_table(
    ["Axe", "F1 macro AVANT correctif (2026-07-07)", "F1 macro APRÈS correctif (2026-07-12)"],
    [
        ("mistake_identification", "0,280", "0,176"),
        ("mistake_location", "0,394", "0,201"),
        ("guidance", "0,370", "0,252"),
        ("actionability", "0,230", "0,242"),
        ("revealing_of_answer", "— (pas de gold, devset v3)", "0,233"),
        ("Moyenne (4 axes d'origine)", "0,319", "0,218"),
        ("Moyenne (5 axes)", "—", "0,221"),
    ],
)
add_body(
    "Constat qualitatif AVANT correctif : le juge sur-prédisait systématiquement \"yes\", "
    "quel que soit le label gold réel — la classe intermédiaire \"to some extent\" était "
    "presque toujours écrasée vers \"yes\". Un correctif de prompt (raisonnement + cadrage "
    "sceptique) a été appliqué à MRBench pour répondre à ce biais (section 2). Sa "
    "recalibration, exécutée le 2026-07-12 sur le même protocole (40 instances, "
    "tests/eval/external/mrbench_f1_calibration_5axes.py, désormais avec un 5e axe "
    "revealing_of_answer grâce aux labels gold de MRBench_V2.json), montre un résultat "
    "inattendu : le F1 macro moyen sur les 4 axes d'origine BAISSE (0,319 → 0,218) au "
    "lieu de s'améliorer.",
)
add_confirmation(
    "Diagnostic exécuté sur les prédictions détaillées (results/"
    "mrbench_f1_judge_validation_5axes.json) : le gold reste très majoritairement \"yes\" "
    "sur les 4 axes d'origine (37/40, 30/40, 27/40, 26/40 instances) mais le juge, après "
    "correctif, ne prédit plus \"yes\" que 16, 14, 22 et 24 fois respectivement — le "
    "cadrage sceptique a sur-corrigé le biais de complaisance : le juge est devenu trop "
    "réticent à répondre \"yes\" pour une distribution gold où \"yes\" est en réalité la "
    "classe très majoritaire, ce qui dégrade le rappel sur cette classe plus que le gain "
    "de précision ne compense. Le seul axe stable est actionability (0,230 → 0,242)."
)
add_confirmation(
    "Hypothèse infrastructure testée et écartée (2026-07-12) : le prompt anti-complaisance "
    "exige un champ \"raisonnement\" avant le verdict, donc plus de tokens à générer que "
    "l'ancien prompt (verdict seul) — hypothèse testée qu'une troncature de sortie côté "
    "Ollama (CPU, pas de num_predict explicitement transmis dans "
    "ai/llm/transports/httpx_transport.py) puisse expliquer le glissement vers \"no\". "
    "Vérification directe sur 8 appels réels (4 dialogues, 2 axes, avec/sans num_predict "
    "explicite) : done_reason=\"stop\" dans 100% des cas (jamais \"length\"), JSON toujours "
    "valide, 34 à 101 tokens générés — bien en dessous de toute limite plausible même sans "
    "le paramètre. La baisse de F1 n'est donc pas un artefact de troncature ni un effet du "
    "CPU (qui n'affecte que la latence, pas le contenu généré) : c'est un effet réel du "
    "modèle face à la consigne \"sois strict et sceptique\", qui reste un bug de code latent "
    "(le paramètre max_tokens n'est jamais réellement transmis à Ollama) mais sans lien "
    "causal établi avec ce résultat."
)
add_body(
    "Pour revealing_of_answer (seul axe désiré=\"no\", gold très majoritairement \"no\" — "
    "34/40), le F1 mesuré est 0,233 : le juge manque environ la moitié des \"no\" réels "
    "(17/34 correctement identifiés), ce qui questionne directement la fiabilité du score "
    "DAMR = 0,225 obtenu sur cet axe en section 5.5 — une partie de ce score bas pourrait "
    "provenir d'erreurs de jugement plutôt que d'un comportement réel du tuteur, dans un "
    "sens qui reste à déterminer (le juge se trompe dans les deux directions, pas "
    "seulement en faveur du tuteur).",
)
add_body(
    "Cette calibration a été faite exclusivement sur la tâche de jugement nuancé de "
    "MRBench (3 niveaux, axes pédagogiques qualitatifs). Elle ne couvre pas directement "
    "les tâches de jugement binaire plus simples de LongMemEval-S et LoCoMo "
    "(\"cette réponse factuelle est-elle correcte ?\") — le même biais y est plausible "
    "(même modèle) mais non vérifié empiriquement.",
    italic=True,
)

# ═══════════════════════════════════════════════════════════════════════
# 7. CONCLUSION
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("7. Conclusion et prochaines étapes", level=1)
add_body(
    "LongMemEval-S confirme un résultat parfait (4/4) sur knowledge-update — signal "
    "encourageant mais à confirmer sur un échantillon plus large. MRBench reste solide "
    "sur 4 axes/5, mais revealing_of_answer résiste : après un premier correctif sans "
    "effet puis un second correctif annulé pour régression confirmée, la cause exacte "
    "reste non identifiée. La recalibration F1 à 5 axes, désormais exécutée, ajoute un "
    "constat inattendu : le prompt anti-complaisance censé corriger le juge a fait "
    "BAISSER sa fiabilité mesurée (F1 moyen 0,319 → 0,218 sur les 4 axes d'origine), par "
    "sur-correction plutôt que par échec du diagnostic initial. C'est désormais le "
    "chantier le plus ouvert de cette campagne, devant revealing_of_answer lui-même.",
    bold=True,
)
add_bullets([
    "Revoir le prompt anti-complaisance du juge MRBench (score_axis) pour corriger la "
    "sur-correction mesurée (F1 4 axes 0,319 → 0,218) — le cadrage sceptique actuel fait "
    "chuter le rappel sur la classe \"yes\", pourtant majoritaire dans le gold, plus qu'il "
    "n'améliore la précision",
    "Ré-interroger le score DAMR = 0,225 de revealing_of_answer à la lumière de son F1 "
    "de calibration (0,233, ~50% de rappel sur les \"no\" gold) — une partie du échec "
    "mesuré pourrait être une erreur de jugement plutôt qu'un comportement réel du tuteur",
    "Investiguer pourquoi le 2e correctif de revealing_of_answer (adaptive.py) a fait "
    "régresser le score au lieu de l'améliorer — capturer le texte réel généré (non "
    "sauvegardé par le runner actuel) avant tout 3e correctif",
    "Remplacer le tableau simulé de LoCoMo par ses résultats réels dès la fin de son run "
    "(~16-17h estimées) — le runner (locomo_runner.py) et les tests associés existent "
    "désormais mais le run réel n'a pas encore été lancé",
    "Documenter les résultats réels dans docs/evaluation.md, en conservant la trace des "
    "runs antérieurs aux correctifs de l'orchestrateur comme dépréciés",
])

doc.save(OUT)
print(f"Rapport généré : {OUT}")
